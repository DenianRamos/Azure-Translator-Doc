pip install requests python-docx

import requests
from docx import Document
import os

api_key = "54tJ8WHmQIQJHrHefOj42HK7vcSmWsNaYfd07BHwKWD89noSEvXFJQQJ99AJACYeBjFXJ3w3AAAbACOGjPyz"
api_endpoint = 'https://api.cognitive.microsofttranslator.com/'
region = "eastus"
target_lang = 'pt-br'

def translate_text(text, target_lang):
    path = '/translate'
    url = api_endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': api_key,
        'Ocp-Apim-Subscription-Region': region,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [{'text': text}]
    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': target_lang
    }

    response = requests.post(url, params=params, headers=headers, json=body)
    return response.json()[0]["translations"][0]["text"]


def translate_doc(file_path):
    doc = Document(file_path)
    translated_paragraphs = []
    for paragraph in doc.paragraphs:
        translated_text = translate_text(paragraph.text, target_lang)
        translated_paragraphs.append(translated_text)
    
    translated_doc = Document()
    for line in translated_paragraphs:
        translated_doc.add_paragraph(line)
    
    translated_file_path = file_path.replace(".docx", f"_{target_lang}.docx")
    translated_doc.save(translated_file_path)
    return translated_file_path

input_file_path = "example.docx"
translate_doc(input_file_path)